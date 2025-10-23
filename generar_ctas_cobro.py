import re
from docx import Document
from num2words import num2words
import os
from PyPDF2 import PdfReader, PdfMerger
from datetime import date
import locale
import re
import pdfplumber

# ================= CONFIGURACIÓN =================
pdf_informe = "INFORME CTAS COBRO HASTA 13 OCTUBRE GOA.pdf"
plantilla_word = "0. FORMATO CARTAS CUENTAS DE COBRO GOA 2024.docx"
carpeta_salida = "ctas_cobro_generadas"

os.makedirs(carpeta_salida, exist_ok=True)

# ================= FUNCIÓN: extraer datos del PDF =================
def extraer_datos(pdf_path):
    print("📖 Extrayendo texto del PDF con pdfplumber...")

    texto = ""
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                texto += page_text + "\n"
            print(f"✅ Página {i}/{len(pdf.pages)} leída")

    # Normalizar texto
    texto = texto.replace("\r", "\n")
    texto = re.sub(r"[ ]{2,}", " ", texto)

    print("\n🔍 Buscando personas, cédulas y totales...")

    # Patrón robusto que detecta cada bloque de persona
    patron = re.compile(
        r"(?P<nombre>[A-ZÁÉÍÓÚÑ ]{5,})\s+Identificaci[oó]n:\s*(?P<cedula>\d{6,12}).*?"
        r"Total para Cuentas por Cobrar[^\d]*(?P<subtotal>\d[\d\.,]*)\s*"
        r"Total:[^\d]*(?P<total>\d[\d\.,]*)",
        re.DOTALL
    )

    personas = []
    for match in patron.finditer(texto):
        nombre = match.group("nombre").strip()
        cedula = match.group("cedula").strip()
        subtotal = match.group("subtotal").replace(".", "").replace(",", ".")
        total = match.group("total").replace(".", "").replace(",", ".")

        personas.append({
            "nombre": nombre,
            "cedula": cedula,
            "subtotal": float(subtotal),
            "total": float(total)
        })

    print(f"\n✔ Se encontraron {len(personas)} personas en el informe.\n")
    for p in personas:
        print(p)

    return personas


# ================= FUNCIÓN: convertir número a texto =================
def numero_a_letras(valor):
    texto = num2words(int(valor), lang='es')
    return texto.replace("y uno", "uno").upper() + " PESOS"

# ================= FUNCIÓN: generar cuenta de cobro =================
def generar_cuenta(persona, plantilla, carpeta_destino):
    doc = Document(plantilla)

    # Fecha actual formateada en español
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except:
        try:
            locale.setlocale(locale.LC_TIME, 'es_CO.UTF-8')
        except:
            pass  # en caso de que el sistema no tenga el locale configurado

    fecha_hoy = date.today().strftime("Bogotá, %d de %B de %Y").capitalize()
    valor_letras = numero_a_letras(persona["total"])
    valor_num = f"{int(persona['total']):,}".replace(",", ".")
    subtotal_num = f"{int(persona['subtotal']):,}".replace(",", ".")

    # Reemplazar textos dentro de los párrafos
    for p in doc.paragraphs:
        texto = p.text
        texto = texto.replace("BOGOTA, 26 DE DICIEMBRE DE 2024", fecha_hoy)
        texto = texto.replace("ALEX ANDRES CARDONA PINILLA", persona["nombre"])
        texto = texto.replace("1.010.183.315", persona["cedula"])
        texto = texto.replace("170.000", subtotal_num)
        texto = texto.replace("Ciento setenta mil pesos".upper(), valor_letras)
        p.text = texto

    # Reemplazar dentro de tablas (si existen)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texto = cell.text
                texto = texto.replace("ALEX ANDRES CARDONA PINILLA", persona["nombre"])
                texto = texto.replace("1.010.183.315", persona["cedula"])
                texto = texto.replace("170.000", subtotal_num)
                texto = texto.replace("Ciento setenta mil pesos".upper(), valor_letras)
                cell.text = texto

    # Guardar Word individual
    doc_path = os.path.join(carpeta_destino, f"{persona['nombre']}.docx")
    doc.save(doc_path)
    return doc_path

# ================= FUNCIÓN: convertir DOCX a PDF =================
def convertir_a_pdf(carpeta):
    pdf_paths = []
    print("📄 Convirtiendo documentos Word a PDF...")
    for archivo in os.listdir(carpeta):
        if archivo.endswith(".docx"):
            input_path = os.path.join(carpeta, archivo)
            os.system(f'libreoffice --headless --convert-to pdf "{input_path}" --outdir "{carpeta}"')
            pdf_paths.append(input_path.replace(".docx", ".pdf"))
            print(f"✅ Convertido: {archivo}")
    return pdf_paths

# ================= FUNCIÓN: unir todos los PDFs =================
def unir_pdfs(lista_pdfs, salida):
    print("\n📚 Unificando todos los PDF generados...")
    merger = PdfMerger()
    for pdf in lista_pdfs:
        merger.append(pdf)
    merger.write(salida)
    merger.close()
    print(f"✅ Archivo final creado: {salida}")

# ================= PROCESO PRINCIPAL =================
if __name__ == "__main__":
    print("🚀 Iniciando generación de cuentas de cobro...\n")
    personas = extraer_datos(pdf_informe)

    print("🧾 Generando cuentas de cobro personalizadas...")
    for i, persona in enumerate(personas, start=1):
        generar_cuenta(persona, plantilla_word, carpeta_salida)
        print(f"✔ Cuenta de cobro creada ({i}/{len(personas)}): {persona['nombre']}")

    pdfs_generados = convertir_a_pdf(carpeta_salida)
    unir_pdfs(pdfs_generados, "Cuentas_Cobro_Consolidadas.pdf")

    print("\n🎉 Proceso completado correctamente.")
